from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

def set_font(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"ℹ️  {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x8A)
    return p

def screenshot(label):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"[ СКРИН: {label} ]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.italic = True
    return p

def bullet(text, bold_part=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part and text.startswith(bold_part):
        r1 = p.add_run(bold_part)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text[len(bold_part):])
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E1F2")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()


# ── Title ────────────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Руководство по работе с ботом поддержки")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

doc.add_paragraph()

# ── Intro ────────────────────────────────────────────────────────────────────

heading("Общее описание", level=1)
body(
    "Бот автоматически создаёт тикеты из двух источников: комментарии к постам канала "
    "и личные сообщения в директ сообщества. Каждый тикет получает уникальный код вида "
    "D2104-01 (D/N — дневная/ночная смена, дата, порядковый номер). "
    "Все тикеты фиксируются в Google Таблице в реальном времени."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — COMMENT TICKET
# ══════════════════════════════════════════════════════════════════════════════

heading("Раздел 1. Тикет из комментария", level=1)

# 1.1
heading("1.1 Как создаётся тикет", level=2)
body(
    "Когда пользователь оставляет комментарий под постом канала, бот автоматически "
    "создаёт тикет в теме «Комментарии» группы поддержки."
)
screenshot("сообщение о новом тикете в группе поддержки — код, имя пользователя, источник «💬 Комментарий», текст и кнопка «⚡ Отреагировать»")

body("Сообщение содержит:")
bullet("Код тикета — уникальный идентификатор", "Код тикета")
bullet("Дата и время обращения", "Дата и время")
bullet("Имя пользователя или @username", "Имя пользователя")
bullet("Источник — 💬 Комментарий", "Источник")
bullet("Текст сообщения или метку медиафайла (🖼 Фото, 🎤 Голосовое и т.д.)", "Текст сообщения")

# 1.2
heading("1.2 Первичная реакция", level=2)
body("Нажми кнопку «⚡ Отреагировать» — это подтверждает, что ты видел обращение.")
screenshot("нажатие кнопки «Отреагировать» — кнопка меняется на «✅ Отреагировано»")
body("После нажатия:")
bullet("Кнопка меняется на «✅ Отреагировано»")
bullet("Пользователь получает уведомление в комментарии о том, что запрос принят")
bullet("В Google Таблице заполняется колонка «Первичная реакция»")
screenshot("уведомление пользователя в треде комментариев")

# 1.3
heading("1.3 Ответ пользователю", level=2)
body("Чтобы ответить пользователю, нажми Reply (Ответить) на любое сообщение тикета в группе поддержки и напиши текст.")
screenshot("администратор делает Reply на сообщение тикета и пишет ответ")
body("Бот автоматически:")
bullet("Доставит ответ пользователю в тред комментариев")
bullet("Опубликует в группе поддержки сообщение «✅ Ответ доставлен» с кнопками")
screenshot("сообщение «Ответ доставлен» с кнопками «Закрыть тикет» и «Удалить ответ пользователю»")
note("Поддерживаемые типы контента: текст, фото, видео, голосовое, кружочек, стикер, файл, GIF.")

# 1.4
heading("1.4 Удаление ошибочного ответа", level=2)
body("Если ответ был отправлен с ошибкой — нажми «🗑 Удалить ответ пользователю».")
screenshot("нажатие кнопки удаления — кнопка меняется на «✅ Удалено»")
body("После нажатия:")
bullet("Сообщение бота в треде комментариев будет удалено")
bullet("Кнопка в группе поддержки меняется на «✅ Удалено»")
bullet("Можно отправить новый ответ, снова ответив на сообщение тикета")

# 1.5
heading("1.5 Закрытие тикета", level=2)
body("Когда вопрос решён — нажми «✅ Закрыть тикет».")
screenshot("кнопка меняется на «🔒 Тикет закрыт»")
body("После закрытия:")
bullet("В Google Таблице заполняется колонка «Время закрытия»")
bullet("Тикет больше не принимает ответы")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DIRECT TICKET
# ══════════════════════════════════════════════════════════════════════════════

heading("Раздел 2. Тикет из директ сообщества", level=1)

# 2.1
heading("2.1 Как создаётся тикет", level=2)
body(
    "Когда пользователь пишет в личные сообщения сообщества (директ), бот создаёт тикет "
    "в теме «Директ» группы поддержки."
)
screenshot("сообщение о новом тикете — источник «📩 Директ», кнопки «⚡ Отреагировать» и «💬 Ответить в директ»")

# 2.2
heading("2.2 Первичная реакция", level=2)
body("Работает так же, как для комментариев — нажми «⚡ Отреагировать».")
screenshot("кнопка меняется на «✅ Отреагировано», пользователь получает уведомление")

# 2.3
heading("2.3 Ответ пользователю", level=2)
body("Для директ-тикетов ответ отправляется вручную через сообщество:")
p = doc.add_paragraph(style="List Number")
r = p.add_run("Нажми кнопку «💬 Ответить в директ»")
r.bold = True
r.font.size = Pt(11)
p.add_run(" — откроется личный чат с пользователем в директ сообщества.").font.size = Pt(11)
p2 = doc.add_paragraph(style="List Number")
p2.add_run("Напиши ответ в этом чате.").font.size = Pt(11)

screenshot("кнопка «Ответить в директ» и открытый чат с пользователем")
body("Бот автоматически обнаружит ответ и опубликует «✅ Ответ доставлен» в группе поддержки.")
screenshot("сообщение «Ответ доставлен» с кнопкой «Закрыть тикет»")
note("Отвечать нужно именно через кнопку «Ответить в директ», а не открывая чат вручную — иначе бот не зафиксирует ответ.")

# 2.4
heading("2.4 Закрытие тикета", level=2)
body("Нажми «✅ Закрыть тикет» после решения вопроса.")
screenshot("кнопка меняется на «🔒 Тикет закрыт»")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — GOOGLE SHEETS
# ══════════════════════════════════════════════════════════════════════════════

heading("Раздел 3. Google Таблица", level=1)
body("Таблица заполняется автоматически в реальном времени по мере работы с тикетом:")

add_table(
    ["Момент", "Что заполняется"],
    [
        ["Создание тикета", "Код, источник, время обращения, сообщение"],
        ["Нажатие «Отреагировать»", "Первичная реакция"],
        ["Отправка ответа", "Вторичная реакция"],
        ["Закрытие тикета", "Время закрытия"],
    ],
)

screenshot("пример заполненной строки в Google Таблице")

body("Колонки «Telegram Chat» и «Telegram Direct» подсвечиваются цветом:")
bullet("🟢 Зелёный — активный источник тикета")
bullet("🔴 Красный — неактивный источник")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — WARNINGS TOPIC
# ══════════════════════════════════════════════════════════════════════════════

heading("Раздел 4. Топик «Предупреждения»", level=1)
body(
    "В группе поддержки есть отдельная служебная тема «Предупреждения». "
    "Она нужна для контроля зависших тикетов и быстрого просмотра текущей очереди."
)
screenshot("тема «Предупреждения» в группе поддержки")

heading("4.1 Когда приходят предупреждения", level=2)
body("Бот отправляет предупреждение, если тикет слишком долго находится без следующего действия:")
bullet("тикет ожидает первичной реакции — нужно нажать «⚡ Отреагировать»")
bullet("тикет ожидает вторичной реакции — нужно ответить пользователю")
bullet("тикет ожидает закрытия — нужно нажать «✅ Закрыть тикет»")
body(
    "Если тикет не закрыт и его статус не меняется, предупреждение будет повторяться "
    "каждые 30 минут, пока ответственный сотрудник не выполнит нужное действие."
)
screenshot("предупреждение о зависшем тикете с указанием номера тикета и текущего статуса")

heading("4.2 Команда /tickets", level=2)
body(
    "В теме «Предупреждения» можно написать команду /tickets. "
    "Бот пришлёт список всех активных тикетов и их текущих статусов."
)
body("Возможные статусы в списке:")
bullet("ожидает первичной реакции")
bullet("ожидает вторичной реакции")
bullet("ожидает закрытия")
note(
    "Если список пустой, значит активных тикетов нет. Закрытые тикеты не отображаются "
    "и удаляются из оперативной базы после синхронизации."
)

doc.add_paragraph()

# ── Final critical warning ───────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(12)
r = p.add_run(
    "ВАЖНО: ВСЕ ТИКЕТЫ ДОЛЖНЫ ПРОХОДИТЬ ПОЛНЫЙ ЦИКЛ: "
    "ОТРЕАГИРОВАНО → РЕАКЦИЯ → ЗАКРЫТИЕ ТИКЕТА"
)
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Документация актуальна на май 2026.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── Save ─────────────────────────────────────────────────────────────────────

out = "support_bot_guide.docx"
doc.save(out)
print(f"Saved: {out}")
